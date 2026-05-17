"""
Smart IP — Teacher Moderation & Paid-Marker Drive document generator.

Produces editable Word (.docx) packs people can complete:
  - 00-Smart-IP-Moderation-Drive-Pack.docx   (brief + every board, one file —
        now also includes, per board: rubric & exemplar summary, marker SOP,
        contract & NDA template, source & consent matrix, calibration guide,
        and a payment & throughput tracker)
  - <Board>-Moderation-Worksheet.docx        (one fillable sheet per board,
        carrying the same paid-marker sections so a marker has everything in
        one file)

Run:  py business-docs/smart-ip/moderation-drive/_generate_docs.py
Regenerate / extend freely — this is the single source of truth for the pack.
*.docx is gitignored repo-wide — only this .py is committed.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent
ROWS_PER_SHEET = 25
TARGET_PER_BOARD = 150

# Engagement / commercial constants (edit here — single source of truth).
COMPANY = "The English Hub"
COMPANY_LEGAL = "The English Hub (Upskill Energy Ltd, trading as The English Hub)"
PRODUCT = "The English Hub AI marker"
PAY_PER_SCRIPT_EXAMPLE = "£1.50"          # illustrative default rate (per approved script)
PAY_GOLD_NOTE = ("Gold/calibration scripts are paid at the same per-script "
                 "rate; they are not optional and not bonus-rated.")
TARGET_GOLD_ACCURACY = "≥ 90% within 1 mark (and ≥ 75% exact) vs. the gold key"
GOLD_FREQUENCY = "roughly 1 in 12 scripts is a blind gold script"

BOARDS = [
    {
        "file": "AQA",
        "name": "AQA",
        "quals": "GCSE English Language (8700)  ·  GCSE English Literature (8702)",
        "scale": "GCSE 9–1 (plus U)",
        "papers": [
            "Eng Language Paper 1 (Explorations in creative reading & writing)",
            "Eng Language Paper 2 (Writers' viewpoints & perspectives)",
            "Eng Literature Paper 1 (Shakespeare & the 19th-century novel)",
            "Eng Literature Paper 2 (Modern texts & poetry)",
        ],
    },
    {
        "file": "Pearson-Edexcel",
        "name": "Pearson Edexcel",
        "quals": "GCSE English Language (1EN0) · English Literature (1ET0)  ·  "
        "IGCSE English Language A (4EA1) · Literature (4ET1)",
        "scale": "GCSE/IGCSE 9–1 (plus U)",
        "papers": [
            "Eng Language Paper 1 (Fiction & imaginative writing)",
            "Eng Language Paper 2 (Non-fiction & transactional writing)",
            "Eng Literature Paper 1 (Shakespeare & post-1914 literature)",
            "Eng Literature Paper 2 (19th-century novel & poetry)",
        ],
    },
    {
        "file": "OCR",
        "name": "OCR",
        "quals": "GCSE English Language (J351)  ·  GCSE English Literature (J352)",
        "scale": "GCSE 9–1 (plus U)",
        "papers": [
            "Eng Language Component 1 (Communicating information & ideas)",
            "Eng Language Component 2 (Exploring effects & impact)",
            "Eng Literature Component 1 (Exploring modern & literary heritage texts)",
            "Eng Literature Component 2 (Exploring poetry & Shakespeare)",
        ],
    },
    {
        "file": "Eduqas-WJEC",
        "name": "Eduqas / WJEC",
        "quals": "GCSE English Language (C700) · English Literature (C720) "
        "(Eduqas; WJEC route equivalents)",
        "scale": "GCSE 9–1 (plus U)",
        "papers": [
            "Eng Language Component 1 (20th-century narrative & creative)",
            "Eng Language Component 2 (19th/21st-century non-fiction & transactional)",
            "Eng Literature Component 1 (Shakespeare & poetry)",
            "Eng Literature Component 2 (Post-1914 prose/drama, 19th-c novel, unseen poetry)",
        ],
    },
    {
        "file": "Cambridge-IGCSE-0500",
        "name": "Cambridge IGCSE 0500 (First Language English)",
        "quals": "Cambridge IGCSE First Language English (0500)",
        "scale": "IGCSE 9–1 / A*–G (record as the centre reports)",
        "papers": [
            "Paper 1 — Reading",
            "Paper 2 — Directed Writing & Composition",
            "Component 3/4 — Coursework / Speaking (where applicable)",
        ],
    },
    {
        "file": "Cambridge-IGCSE-0990",
        "name": "Cambridge IGCSE 0990 (English Language 9–1)",
        "quals": "Cambridge IGCSE English Language – 9–1 (0990)",
        "scale": "IGCSE 9–1 (plus U)",
        "papers": [
            "Paper 1 — Reading",
            "Paper 2 — Directed Writing & Composition",
            "Component 3/4 — Coursework / Speaking (where applicable)",
        ],
    },
]

QTYPES = [
    "Reading – language analysis",
    "Reading – structure analysis",
    "Reading – evaluation / critical response",
    "Reading – summary / synthesis",
    "Reading – comparison",
    "Writing – creative / descriptive / narrative",
    "Writing – transactional / persuasive / discursive",
    "Literature – extract-based essay",
    "Literature – whole-text essay",
    "Literature – poetry (anthology / comparison)",
    "Literature – unseen poetry",
]
BANDS = ["U", "1–3", "4–5", "6–7", "8–9"]

# ── Per-board rubric & exemplar summary ──────────────────────────────────────
# Concise, accurate distillation of the published assessment-objective intent
# (NOT a reproduction of any board's copyrighted mark scheme — descriptors are
# paraphrased to the AO's purpose so markers calibrate consistently). Keyed by
# the board "file" id. `aos` = the assessment objectives that dominate; `bands`
# = what each grade band looks like in practice; `exemplars` = short worked
# illustrations of the marking decision (answer sketch → mark → why).
GENERIC_BANDS = [
    ("8–9", "Perceptive, conceptualised. Analysis is judicious; precise, "
            "well-chosen textual detail integrated into a convincing line of "
            "argument; sophisticated, accurate writing where assessed."),
    ("6–7", "Clear, consistent understanding. Explained analysis with apt "
            "references; coherent structure; accurate, controlled expression."),
    ("4–5", "Some understanding, mostly relevant. Comments rather than "
            "analysis; references support but aren't probed; generally clear "
            "writing with lapses."),
    ("1–3", "Simple, limited. Largely descriptive/retelling; few or "
            "undeveloped references; frequent technical errors."),
    ("U",   "Nothing creditable, or no legible/relevant response to the task."),
]

RUBRIC = {
    "AQA": {
        "aos": [
            "AO1 — identify & interpret explicit/implicit information; "
            "select & synthesise (Lang P2).",
            "AO2 — analyse language/structure/form; subject terminology.",
            "AO3 — compare writers' ideas & perspectives (Lang P2).",
            "AO4 — evaluate texts critically with textual reference (Lang P1 Q4).",
            "AO5/AO6 — content & organisation / technical accuracy (Lang "
            "writing). Lit: AO1 response & references, AO2 methods, AO3 "
            "context, AO4 vocabulary/spelling/punctuation.",
        ],
        "bands": GENERIC_BANDS,
        "exemplars": [
            ("Lang P1 Q2 (language analysis, 8 marks): answer names a simile "
             "and a verb choice and explains the effect on the reader with "
             "precise quotation, but does not explore connotations deeply.",
             "5/8 (top of mid band)",
             "Clear analysis with apt terminology = secure mid; not "
             "perceptive/whole-quotation exploratory, so not top band. AI gave "
             "7 — over-credited; correction reason: 'depth of language "
             "analysis below top band — explanation not exploratory'."),
            ("Lang P1 Q5 (creative writing, 24+16): vivid controlled "
             "description, varied sentences, ambitious vocabulary, a few "
             "punctuation slips.",
             "Content 19/24, Technical 13/16",
             "Compelling & crafted = top of content; mostly accurate with "
             "ambitious range = high technical band. AI's 16/24 under-marked "
             "content; reason: 'crafting & vocabulary range merit upper band'."),
            ("Lit (Macbeth extract->whole): thoughtful thesis on guilt, "
             "well-chosen quotations analysed for method, some context.",
             "Upper Level 5 of 6",
             "Thoughtful/developed with method focus = Level 5; not "
             "consistently exploratory/critical = not Level 6."),
        ],
    },
    "Pearson-Edexcel": {
        "aos": [
            "AO1 — identify/interpret; maintain a critical style (Lit).",
            "AO2 — analyse language, form, structure; terminology.",
            "AO3 — compare; (Lit) relate texts to contexts.",
            "AO4 — evaluate (Lang); (Lit) accuracy where assessed.",
            "AO5/AO6 — communication & organisation / spelling, punctuation, "
            "grammar (Lang writing).",
        ],
        "bands": GENERIC_BANDS,
        "exemplars": [
            ("Lang P1 imaginative writing: well-structured narrative, "
             "deliberate devices, controlled tone, accurate.",
             "Level 4 (upper) communication & language",
             "Effective, sustained crafting = Level 4; isolated lapses keep "
             "it off the very top. AI placed Level 3 — reason: "
             "'structural control & device choice are sustained, not just "
             "appropriate'."),
            ("Lang P2 non-fiction comparison (AO3): identifies shared theme, "
             "compares two clear methods with quotation, link explained.",
             "Mid band",
             "Clear comparison of methods = mid; not perceptive synthesis of "
             "writers' perspectives = not top."),
            ("Lit post-1914 essay: clear argument, integrated references, "
             "some exploration of dramatic method and context.",
             "Mid–upper band",
             "Explained, well-supported response; not consistently analytical "
             "across method+context for top band."),
        ],
    },
    "OCR": {
        "aos": [
            "AO1 — identify & interpret; communicate clearly (Lang).",
            "AO2 — analyse language/structure; terminology.",
            "AO3 — compare ideas/perspectives & how conveyed (Lang); (Lit) "
            "context.",
            "AO4 — evaluate / (Lang) technical accuracy in writing.",
            "Lit AO1–AO4 — informed personal response, methods, context, "
            "accuracy.",
        ],
        "bands": GENERIC_BANDS,
        "exemplars": [
            ("Lang C2 (effects & impact): explains writer's structural choice "
             "and one language pattern with quotation; effect on reader "
             "stated, not probed.",
             "Mid band (Level 4 of 6-style)",
             "Explained analysis = mid; not sustained/perceptive = not top. "
             "AI over-marked at top — reason: 'effect asserted not analysed "
             "in depth'."),
            ("Lang creative: controlled, well-organised, ambitious "
             "vocabulary mostly accurate.",
             "Upper band content; secure technical",
             "Crafted & organised = upper; technical range secure with minor "
             "slips."),
            ("Lit poetry comparison: clear comparative argument, apt "
             "quotations, some method analysis.",
             "Mid–upper band",
             "Developed comparison; not consistently exploratory of method "
             "for the very top."),
        ],
    },
    "Eduqas-WJEC": {
        "aos": [
            "AO1 — identify/interpret; select & synthesise.",
            "AO2 — analyse language/structure/form; terminology.",
            "AO3 — compare writers' ideas & perspectives.",
            "AO4 — evaluate critically (Lang reading).",
            "AO5/AO6 — communicating & organising / accuracy (Lang writing). "
            "Lit AO1–AO4 as published.",
        ],
        "bands": GENERIC_BANDS,
        "exemplars": [
            ("Lang C1 narrative: engaging, structured, varied sentence "
             "forms, ambitious punctuation mostly secure.",
             "Band 4–5 of 5 (upper)",
             "Engaging & controlled crafting = upper band; minor accuracy "
             "slips keep off absolute top. AI gave mid — reason: 'sentence "
             "variety & structural shaping are sustained'."),
            ("Lang C2 (19th/21st non-fiction synthesis): synthesises both "
             "sources accurately, clear inference.",
             "Upper band AO1",
             "Accurate synthesis with inference = upper; full perceptive "
             "synthesis needed for the very top."),
            ("Lit unseen poetry: clear personal response, analyses imagery "
             "and one structural feature with quotation.",
             "Mid–upper band",
             "Explained analysis of method; not consistently perceptive for "
             "top band."),
        ],
    },
    "Cambridge-IGCSE-0500": {
        "aos": [
            "Reading (R1–R5) — understand explicit/implicit meaning & "
            "attitudes; analyse how writers achieve effects; select & "
            "synthesise.",
            "Writing (W1–W5) — articulate experience/ideas; organise & "
            "sequence; use a range of accurate vocabulary, sentence "
            "structures, spelling & punctuation; appropriate register.",
        ],
        "bands": [
            ("8–9", "Reading: sophisticated, well-selected detail; analysis "
                    "of writer's effects. Writing: highly effective, "
                    "consistently accurate, strong register control."),
            ("6–7", "Reading: clear understanding incl. some implicit; "
                    "relevant detail. Writing: effective, well-organised, "
                    "generally accurate."),
            ("4–5", "Reading: general understanding; some relevant detail. "
                    "Writing: mostly clear & organised; lapses in accuracy."),
            ("1–3", "Reading: literal/limited. Writing: simple, frequent "
                    "errors that may impede meaning."),
            ("U",   "No creditable reading response / writing not "
                    "communicable or off-task."),
        ],
        "exemplars": [
            ("P1 directed-reading: selects relevant explicit points and one "
             "implicit idea, mostly in own words.",
             "Upper-mid (R band)",
             "Clear incl. some implicit understanding = upper-mid; not "
             "consistently analytical of writer's effects for top. AI "
             "over-credited — reason: 'implicit understanding present but "
             "not developed'."),
            ("P2 composition (descriptive): vivid, well-sequenced, varied "
             "sentences, accurate, controlled register.",
             "Top band W",
             "Highly effective + consistently accurate + register = top band."),
            ("P2 directed writing: covers task purpose/audience, clear "
             "structure, some persuasive method, minor slips.",
             "Mid–upper band",
             "Effective & organised; full range/precision needed for top."),
        ],
    },
    "Cambridge-IGCSE-0990": {
        "aos": [
            "Reading (R1–R5) — explicit & implicit meaning, attitudes; "
            "analyse writers' effects; select & synthesise.",
            "Writing (W1–W5) — communicate clearly/effectively; organise; "
            "accurate vocabulary, sentence structure, spelling & "
            "punctuation; register for purpose & audience.",
        ],
        "bands": [
            ("8–9", "Reading: perceptive selection & analysis of effects. "
                    "Writing: highly effective, accurate, precise register."),
            ("6–7", "Reading: clear, some implicit. Writing: effective, "
                    "organised, generally accurate."),
            ("4–5", "Reading: general; some detail. Writing: mostly clear; "
                    "accuracy lapses."),
            ("1–3", "Reading: literal/limited. Writing: simple, frequent "
                    "errors."),
            ("U",   "No creditable response / off-task."),
        ],
        "exemplars": [
            ("P1 reading (analysis of effects): identifies two writer "
             "techniques with quotation; effect explained, not fully "
             "analysed.",
             "Mid band",
             "Explained effects = mid; perceptive analysis needed for top. "
             "AI gave top — reason: 'technique effects explained but not "
             "probed'."),
            ("P2 composition (argumentative): clear thesis, organised, "
             "controlled register, accurate.",
             "Upper band",
             "Effective + accurate + register = upper; sustained "
             "sophistication for the very top."),
            ("P2 summary: synthesises required points concisely in own "
             "words, accurate.",
             "Upper band R",
             "Accurate, concise own-words synthesis = upper band."),
        ],
    },
}


def _board_rubric(b):
    """Rubric record for a board, falling back to a generic GCSE/IGCSE
    framing so a newly-added board still produces a usable section."""
    return RUBRIC.get(b["file"], {
        "aos": [
            "AO1 — identify/interpret; (writing) communicate & organise.",
            "AO2 — analyse language/structure/form with terminology.",
            "AO3 — compare ideas/perspectives; (Lit) context.",
            "AO4 — evaluate critically / technical accuracy.",
        ],
        "bands": GENERIC_BANDS,
        "exemplars": [
            ("A clear, well-supported response with explained (not yet "
             "exploratory) analysis.",
             "Secure mid band",
             "Explained + apt references = mid; perceptive conceptualised "
             "treatment needed for top band."),
        ],
    })


TRACKER_COLS = [
    ("#", 0.4),
    ("Paper / Component", 1.7),
    ("Question type", 1.7),
    ("Target band", 0.7),
    ("Anon. script ref", 1.1),
    ("AI mark", 0.6),
    ("Teacher mark", 0.7),
    ("Δ (T−AI)", 0.6),
    ("Adjustment reason — WHY the AI mark was changed (the training signal)", 2.6),
    ("Approved + 'suitable for training' set? (Y/N)", 1.1),
    ("Moderator / date", 1.0),
]


def _landscape_narrow(doc):
    s = doc.sections[-1]
    s.orientation = WD_ORIENT.LANDSCAPE
    w, h = s.page_height, s.page_width
    s.page_width, s.page_height = w, h
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, m, Inches(0.5))


def _hr_box(doc, text, fill="F2F2F2"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    return p


def _set_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = Inches(widths[i])


def add_board_section(doc, b, first=False):
    if not first:
        doc.add_page_break()

    h = doc.add_heading(f"{b['name']} — Teacher Moderation Worksheet", level=1)
    sub = doc.add_paragraph()
    sub.add_run(b["quals"]).italic = True
    doc.add_paragraph(f"Grade scale: {b['scale']}    |    "
                      f"Target: ≥ {TARGET_PER_BOARD} dual-marked, teacher-approved "
                      f"scripts for this board, spread across papers/components.")

    _hr_box(doc, "STEP 0 — CONSENT (THE GATE). No consent → DO NOT approve for "
                 "training. This is the single biggest lever.")
    doc.add_paragraph(
        "For every script you include below, confirm BEFORE approving it for "
        "training:", style="Intense Quote")
    for t in [
        "☐  The student has given AI-processing consent (Settings ▸ Privacy ▸ AI).",
        "☐  Training-use consent is on (B2C: 'AI training opt-in'; B2B: school "
        "data-sharing enabled).",
        "☐  If the student is under 18 — verifiable parental/guardian consent is "
        "on record.",
        "☐  If ANY box is unticked: review for feedback if you wish, but DO NOT "
        "tick 'suitable for training' / do not approve into the corpus.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_paragraph(
        "The mechanic:  corpus growth = (teacher-approved submissions) × "
        "(training toggle ON) × (consent satisfied). Each time you Approve a "
        "submission in /school/marking with 'suitable for training' set, it "
        "becomes eligible and is anonymised into the training corpus.")

    doc.add_heading("How to complete this worksheet", level=2)
    for t in [
        "Genuinely review every script — adjust the mark and feedback where the "
        "AI is wrong. Rubber-stamping produces worthless labels.",
        "ALWAYS fill the 'Adjustment reason' — that correction signal is the "
        "actual training value (it is stored in teacher_moderations).",
        "Spread your scripts across paper/component × question-type × grade band "
        "(use the coverage grid below). Don't let one cohort dominate.",
        f"Aim for ≥ {TARGET_PER_BOARD} dual-marked scripts for {b['name']} "
        "(threshold in evals/datasets/REAL-DATA-PROTOCOL.md for a defensible "
        "accuracy figure). Below that you have a pipeline, not a certifiable "
        "model signal.",
        f"This sheet holds {ROWS_PER_SHEET} rows — run as many copies as needed "
        "to reach the target.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("Collection tracker", level=2)
    headers = [c[0] for c in TRACKER_COLS]
    widths = [c[1] for c in TRACKER_COLS]
    tbl = doc.add_table(rows=1 + ROWS_PER_SHEET, cols=len(headers))
    tbl.style = "Table Grid"
    for i, txt in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(8)
    for r in range(1, 1 + ROWS_PER_SHEET):
        tbl.rows[r].cells[0].text = str(r)
        for i in range(len(headers)):
            for para in tbl.rows[r].cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    _set_widths(tbl, widths)

    doc.add_paragraph()
    doc.add_heading("Coverage grid (tally each script — keep the spread balanced)",
                    level=2)
    cov = doc.add_table(rows=1 + len(QTYPES), cols=1 + len(BANDS))
    cov.style = "Table Grid"
    cov.rows[0].cells[0].text = "Question type \\ Grade band"
    cov.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    for j, band in enumerate(BANDS):
        c = cov.rows[0].cells[1 + j]
        c.text = ""
        rr = c.paragraphs[0].add_run(band)
        rr.bold = True
    for i, qt in enumerate(QTYPES):
        cov.rows[1 + i].cells[0].text = qt
        for para in cov.rows[1 + i].cells[0].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)
    _set_widths(cov, [2.6] + [1.3] * len(BANDS))

    doc.add_paragraph()
    doc.add_heading("Moderator sign-off", level=2)
    for line in [
        "Moderator name: ______________________________________________",
        "Role / examining qualification: ________________________________",
        "Papers/components covered on this sheet: _______________________",
        "Scripts completed on this sheet: ______   Running total for "
        f"{b['name']}: ______ / {TARGET_PER_BOARD}",
        "I confirm consent (Step 0) was verified for EVERY script approved "
        "into training on this sheet.",
        "Signature: ____________________________     Date: ______________",
    ]:
        doc.add_paragraph(line)


def build_brief(doc):
    t = doc.add_heading("Smart IP — Teacher Moderation Drive", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Collecting clean, teacher-approved, anonymised marking data to train "
        "and validate The English Hub's AI marker. Controlled and human-"
        "reviewed — there is no autonomous retraining."
    ).italic = True

    doc.add_heading("Step 0 — Consent (the gate)", level=1)
    doc.add_paragraph(
        "Make AI-processing + training-use consent an explicit checkbox in "
        "student/school onboarding (and verifiable parental consent for "
        "under-18s). Without it, nothing is collectible. This is your single "
        "biggest lever — a record only enters the corpus when the student has "
        "BOTH the privacy flag (B2C 'AI training opt-in' / B2B school data-"
        "sharing) AND an explicit AI-processing consent, plus parental consent "
        "if a minor. The live pipeline enforces this and refuses to write a "
        "training row otherwise.")

    doc.add_heading("The mechanic", level=1)
    doc.add_paragraph(
        "Corpus growth = (teacher-approved submissions) × (training toggle on) "
        "× (consent satisfied). Every time a teacher hits Approve with "
        "'suitable for training' in /school/marking, that submission becomes "
        "eligible; prepareTrainingRecord anonymises it into the corpus.")

    doc.add_heading("Fastest, highest-quality path — a structured moderation "
                    "drive (recommended)", level=1)
    for t in [
        "Recruit 3–6 qualified English teachers/examiners (your own staff or "
        "contractors) with consent in place.",
        "Feed a targeted set through /marking (or /school/marking): spread "
        "across board × paper × question-type × grade band — don't let it skew "
        "to one cohort.",
        "Teachers genuinely review (edit the mark/feedback, fill the adjustment "
        "reason) — rubber-stamping produces worthless labels. The "
        "teacher_moderations table captures the correction signal that is the "
        "actual training value.",
        f"Target ≥ {TARGET_PER_BOARD} dual-marked scripts per board/paper (the "
        "threshold in evals/datasets/REAL-DATA-PROTOCOL.md for a defensible "
        "accuracy figure). Below that you have a pipeline but not a certifiable "
        "model signal.",
    ]:
        doc.add_paragraph(t, style="List Number")

    doc.add_heading("How to run it", level=1)
    for t in [
        "Give each moderator the worksheet for their board (separate .docx "
        "files accompany this pack — one per exam board).",
        "Verify Step 0 consent per script BEFORE ticking 'suitable for "
        "training'.",
        "Track progress in /admin/ai-marking and /admin/model-performance; "
        "export from /admin/training-data when a board hits its target; the "
        "evals/ harness then produces the real accuracy figure.",
        "Boards in this pack: AQA · Pearson Edexcel · OCR · Eduqas/WJEC · "
        "Cambridge IGCSE 0500 · Cambridge IGCSE 0990.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # one-page board target summary
    doc.add_heading("Board targets at a glance", level=1)
    s = doc.add_table(rows=1 + len(BOARDS), cols=4)
    s.style = "Table Grid"
    for i, htxt in enumerate(["Exam board", "Qualifications", "Grade scale",
                              "Target (dual-marked, approved)"]):
        c = s.rows[0].cells[i]
        c.text = ""
        c.paragraphs[0].add_run(htxt).bold = True
    for i, b in enumerate(BOARDS):
        s.rows[1 + i].cells[0].text = b["name"]
        s.rows[1 + i].cells[1].text = b["quals"]
        s.rows[1 + i].cells[2].text = b["scale"]
        s.rows[1 + i].cells[3].text = f"≥ {TARGET_PER_BOARD}"


def _para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def _signature_block(doc, lines):
    for line in lines:
        doc.add_paragraph(line)


# ── New section 1: per-board rubric & exemplar summary ───────────────────────
def add_rubric_summary(doc, b):
    rb = _board_rubric(b)
    doc.add_heading(f"{b['name']} — Rubric & exemplar summary (calibrate to "
                    "THIS before you mark)", level=2)
    _para(doc,
          "Use this to anchor your judgement to the mark-scheme INTENT. It is "
          "a paraphrase of the published assessment objectives and level "
          "descriptors — not a copy of the board's mark scheme. Always defer "
          "to the official mark scheme for the specific question; this is the "
          "shared calibration reference.", italic=True)

    doc.add_heading("Assessment objectives that drive the mark", level=3)
    for ao in rb["aos"]:
        doc.add_paragraph(ao, style="List Bullet")

    doc.add_heading("Grade-band descriptors (what each band looks like)", level=3)
    bt = doc.add_table(rows=1 + len(rb["bands"]), cols=2)
    bt.style = "Table Grid"
    for i, htxt in enumerate(["Band", "What it looks like in practice"]):
        c = bt.rows[0].cells[i]
        c.text = ""
        c.paragraphs[0].add_run(htxt).bold = True
    for i, (band, desc) in enumerate(rb["bands"]):
        bt.rows[1 + i].cells[0].text = band
        bt.rows[1 + i].cells[1].text = desc
        for cidx in (0, 1):
            for para in bt.rows[1 + i].cells[cidx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    _set_widths(bt, [0.8, 8.7])

    doc.add_paragraph()
    doc.add_heading("Worked exemplars (the marking decision, explained)", level=3)
    et = doc.add_table(rows=1 + len(rb["exemplars"]), cols=3)
    et.style = "Table Grid"
    for i, htxt in enumerate(["Answer sketch", "Mark / band",
                              "Why — and the correction signal"]):
        c = et.rows[0].cells[i]
        c.text = ""
        c.paragraphs[0].add_run(htxt).bold = True
    for i, (sketch, mark, why) in enumerate(rb["exemplars"]):
        et.rows[1 + i].cells[0].text = sketch
        et.rows[1 + i].cells[1].text = mark
        et.rows[1 + i].cells[2].text = why
        for cidx in (0, 1, 2):
            for para in et.rows[1 + i].cells[cidx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    _set_widths(et, [3.3, 1.5, 4.7])


# ── New section 2: marker SOP (exact console steps) ──────────────────────────
def add_marker_sop(doc, b):
    doc.add_heading(f"{b['name']} — Marker SOP (exact steps in the console)",
                    level=2)
    _para(doc, "Follow this every script. The value is in genuine correction, "
               "not throughput.", italic=True)

    doc.add_heading("Step-by-step", level=3)
    for i, step in enumerate([
        "Log in at https://www.theenglishhub.app with your marker account "
        "(magic-link / password as issued). You will land authenticated.",
        "Go to /marker — your assigned queue for this board/batch loads "
        "(scripts are assigned to you; you only see your own).",
        "Open the next script. Read the student answer IN FULL before looking "
        "at the AI draft — form your own provisional mark first.",
        "Reveal the AI draft (score, grade band, AO/band breakdown, "
        "feedback). Compare it against your provisional judgement and the "
        "rubric & exemplar summary in this pack.",
        "Correct the mark and the per-AO/band marks where the AI is wrong. "
        "Edit the feedback so it is accurate, specific and student-usable.",
        "Fill the ADJUSTMENT REASON — one or two precise sentences on WHY the "
        "AI mark/feedback was changed (e.g. 'AO2 over-credited: analysis is "
        "explained not exploratory'). This free-text is the single most "
        "valuable training signal — never leave it blank when you changed "
        "anything; if the AI was correct, say so briefly ('AI mark correct; "
        "confirmed against band descriptor').",
        "If the script is unsuitable (off-task, not creditable, wrong paper, "
        "or you cannot mark it confidently) choose Reject / Send back with a "
        "reason — do NOT approve it into the corpus.",
        "Click APPROVE & NEXT. Approve commits your decision, records the "
        "moderation, and (when consent applies, see the matrix) makes the "
        "script eligible for the anonymised training corpus. The next script "
        "loads automatically.",
        "At the end of a session, fill the payment & throughput tracker sheet "
        "(scripts done, gold scripts seen) and submit it as instructed.",
    ], start=1):
        doc.add_paragraph(f"{i}.  {step}", style="List Number")

    doc.add_heading("What 'good' looks like", level=3)
    for g in [
        "You marked the script yourself and the final mark reflects YOUR "
        "judgement against the rubric, not a nudge of the AI's.",
        "Every adjustment has a clear, specific reason a colleague could act "
        "on. Reasons are about the AO/criterion, not vague ('a bit generous' "
        "→ 'AO2: device effects asserted, not analysed').",
        "Feedback is accurate and usable by a student.",
        "Your marks are consistent across the session (gold scripts confirm "
        "this — see the calibration guide).",
    ]:
        doc.add_paragraph(g, style="List Bullet")

    _hr_box(doc, "NO RUBBER-STAMPING. Approving the AI draft unchanged without "
                 "genuinely checking it produces worthless labels, corrupts "
                 "the training signal, and is a breach of the marker "
                 "agreement. Pay is for genuine review, not clicks. Random "
                 "gold scripts and audit re-marks detect rubber-stamping.")


# ── New section 3: marker contract & NDA template ────────────────────────────
def add_contract_nda(doc, b):
    doc.add_heading("Marker Engagement, IP Assignment & NDA — template",
                    level=2)
    _para(doc, f"Between {COMPANY_LEGAL} (the \"Company\") and the individual "
               "named below (the \"Marker\"). TEMPLATE — not legal advice; "
               "have it reviewed by a qualified lawyer before use. Complete "
               "the bracketed fields.", italic=True)

    clauses = [
        ("1. Engagement & status",
         "The Company engages the Marker on an independent-contractor basis "
         "to review, correct and approve AI-generated marking of English "
         "Language/Literature scripts for the boards assigned. Nothing here "
         "creates employment, partnership or agency. The Marker provides "
         "services personally and may not subcontract."),
        ("2. Services & standard",
         "The Marker will genuinely review each assigned script, correct the "
         "mark/feedback per the applicable board mark scheme, and record an "
         "honest adjustment reason, following the Marker SOP and calibration "
         "requirements in this pack. Rubber-stamping, fabricating reasons, or "
         "approving unreviewed output is a material breach."),
        ("3. IP assignment of marking output",
         "The Marker hereby assigns to the Company, with full title "
         "guarantee, all present and future intellectual property rights in "
         "all marking output produced under this engagement — including "
         "corrected marks, grade/AO judgements, written feedback, adjustment "
         "reasons, annotations and any derived labels or data — absolutely "
         "and for all media and territories, for the Company to use without "
         "restriction including to train, validate and improve the Company's "
         "AI marking models and products. Where rights cannot assign by law "
         "the Marker grants an irrevocable, worldwide, royalty-free, "
         "sublicensable licence and waives all moral rights to the fullest "
         "extent permitted. The Marker will sign any further documents the "
         "Company reasonably needs to perfect this assignment."),
        ("4. Confidentiality",
         "The Marker will keep confidential all Company information, "
         "including scripts, student data, mark schemes, prompts, model "
         "behaviour, the existence and detail of gold/calibration scripts, "
         "rates and these terms; will use it only to perform the services; "
         "and will not copy, retain or disclose it. Confidentiality survives "
         "termination indefinitely (and as long as the law requires for "
         "personal data)."),
        ("5. Data protection & handling",
         "Scripts may contain personal data of students, some of them "
         "minors. The Marker acts only on the Company's documented "
         "instructions as a processor: access scripts ONLY inside the "
         "Company console; never download, screenshot, copy, transmit or "
         "store scripts or student data outside the console; never attempt "
         "to identify or contact a student; use a secure, password-protected "
         "device; report any suspected data breach to the Company without "
         "undue delay and within 24 hours. Commissioned/specimen material "
         "carries no pupil personal data (see the Source & consent matrix); "
         "platform scripts are consent-gated by the Company."),
        ("6. Quality, audit & re-marking",
         "The Company may insert blind gold/calibration scripts and re-mark "
         "samples to audit quality. The Marker must maintain the calibration "
         "standard in this pack. Sustained drift or detected rubber-stamping "
         "may lead to retraining, withholding of disputed payment, or "
         "termination."),
        ("7. Payment",
         f"The Company pays the Marker per APPROVED script at the rate set "
         f"out in the engagement letter (illustratively {PAY_PER_SCRIPT_EXAMPLE} "
         "per approved script; the rate in the signed letter governs). An "
         "\"approved script\" is one the Marker has genuinely reviewed and "
         "approved (or approved-with-corrections) in the console and which "
         "passes quality checks. " + PAY_GOLD_NOTE + " Scripts rejected for "
         "rubber-stamping or quality failure are not payable. Payment is "
         "monthly in arrears against the Company's throughput record within "
         "30 days of a correct invoice. The Marker is responsible for their "
         "own tax and NICs."),
        ("8. Independent obligations",
         "The Marker warrants the services are their own work, performed "
         "with reasonable skill and care and the stated examining "
         "competence, and that they are entitled to provide them."),
        ("9. Term & termination",
         "Either party may terminate on 7 days' written notice, or "
         "immediately for material breach. Clauses 3, 4, 5 and accrued "
         "payment obligations survive termination."),
        ("10. Governing law",
         "These terms are governed by the laws of England and Wales and "
         "subject to the exclusive jurisdiction of its courts."),
    ]
    for title, body in clauses:
        doc.add_heading(title, level=3)
        doc.add_paragraph(body)

    doc.add_heading("Signatures", level=3)
    _signature_block(doc, [
        "Marker name: ______________________________________________",
        "Examining qualification / experience: ____________________________",
        "Boards engaged for: ______________________________________________",
        "Agreed rate per approved script: £__________   Effective date: __________",
        "Marker signature: ____________________________   Date: ____________",
        "",
        "For and on behalf of " + COMPANY + ":",
        "Name / role: ______________________________________________",
        "Signature: ____________________________   Date: ____________",
    ])


# ── New section 4: source & consent matrix ───────────────────────────────────
def add_source_consent_matrix(doc, b):
    doc.add_heading("Source & consent matrix (when is consent required?)",
                    level=2)
    _para(doc,
          "Consent is SOURCE-AWARE and enforced in code by "
          "prepareTrainingRecord. Real-pupil work is consent-gated; "
          "commissioned/specimen material has no pupil data subject and is "
          "consent-exempt. The pipeline refuses to write a training row that "
          "fails its applicable gate.", italic=True)

    rows = [
        ("source = 'commissioned'",
         "Paid-marker / author-produced answers. No identifiable pupil.",
         "NOT required (no pupil data subject).",
         "Mark & approve normally. The record is anonymised and enters the "
         "corpus without a consent check (logged as consent-exempt)."),
        ("source = 'specimen'",
         "Exam-board / specimen / past-paper style material. No pupil.",
         "NOT required (no pupil data subject).",
         "Mark & approve normally. Treated as consent-exempt, same as "
         "commissioned."),
        ("source = 'b2c_self'",
         "A real B2C student's own submission.",
         "REQUIRED: student AI-processing consent + 'AI training opt-in' "
         "privacy flag, plus verifiable parental consent if under 18.",
         "Only approve into training if consent is satisfied; otherwise "
         "review for feedback but do NOT approve into the corpus."),
        ("source = 'b2b_class'",
         "A real school/class pupil's submission.",
         "REQUIRED: school data-sharing enabled + pupil AI-processing "
         "consent, plus verifiable parental consent if a minor.",
         "Only approve into training if consent is satisfied; otherwise do "
         "NOT approve into the corpus."),
    ]
    mt = doc.add_table(rows=1 + len(rows), cols=4)
    mt.style = "Table Grid"
    for i, htxt in enumerate(["Source", "What it is",
                              "Pupil consent required?",
                              "What the marker does"]):
        c = mt.rows[0].cells[i]
        c.text = ""
        c.paragraphs[0].add_run(htxt).bold = True
    for i, r in enumerate(rows):
        for cidx in range(4):
            mt.rows[1 + i].cells[cidx].text = r[cidx]
            for para in mt.rows[1 + i].cells[cidx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    _set_widths(mt, [1.6, 2.5, 2.4, 3.0])

    _hr_box(doc, "If you are unsure which source a script is, or whether "
                 "consent applies — STOP and ask the Company. Never approve a "
                 "real-pupil script into training on assumption.")


# ── New section 5: calibration guide ─────────────────────────────────────────
def add_calibration_guide(doc, b):
    doc.add_heading("Calibration guide (gold scripts & staying accurate)",
                    level=2)
    doc.add_paragraph(
        "Gold scripts are pre-marked by a senior examiner to an agreed key. "
        f"They are inserted BLIND into your queue ({GOLD_FREQUENCY}) — you "
        "cannot tell a gold script from a normal one, so mark every script as "
        "if it were gold.")
    doc.add_heading("The standard you must hold", level=3)
    for s in [
        f"Target gold accuracy: {TARGET_GOLD_ACCURACY}.",
        "Per-AO/criterion judgements should match the gold key's rationale, "
        "not just the total.",
        "Consistency matters as much as the single mark — the same script "
        "should get the same mark from you on a different day.",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_heading("What happens if you drift", level=3)
    for s in [
        "Mild drift: you get a calibration note with the gold key and "
        "rationale; re-read the rubric summary and recalibrate.",
        "Sustained drift below target: a short re-training/recalibration "
        "session before you continue marking; affected scripts may be "
        "re-marked.",
        "Persistent drift or detected rubber-stamping: marking is paused; "
        "disputed scripts are not paid; the engagement may be terminated "
        "(see contract clause 6).",
        "Gold scripts themselves are still paid at the normal per-script "
        "rate — they are part of the job, not a penalty.",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_heading("Self-check before you submit a session", level=3)
    for s in [
        "Did I mark each script myself before seeing the AI draft?",
        "Are my adjustment reasons specific and criterion-based?",
        "Would my marks survive a senior examiner re-mark?",
    ]:
        doc.add_paragraph(s, style="List Bullet")


# ── New section 6: payment & throughput tracker sheet ────────────────────────
PAYTRACK_COLS = [
    ("Date", 0.9),
    ("Marker name", 1.4),
    ("Board", 1.2),
    ("Batch / queue ref", 1.3),
    ("Scripts reviewed", 1.0),
    ("of which approved", 1.0),
    ("of which gold (blind)", 1.0),
    ("Rejected / sent back", 1.0),
    ("Rate per approved script (£)", 1.1),
    ("Amount due (£) = approved × rate", 1.4),
    ("Marker signature", 1.4),
]


def add_payment_tracker(doc, b):
    doc.add_heading(f"{b['name']} — Payment & throughput tracker", level=2)
    _para(doc,
          "Fill one row per marking session and submit as instructed. The "
          "Company reconciles this against the system throughput record "
          "(/admin/marker-pay). Pay = approved scripts × agreed rate; gold "
          "scripts are included in 'approved' and paid normally; "
          "rubber-stamped or quality-failed scripts are not payable.",
          italic=True)
    headers = [c[0] for c in PAYTRACK_COLS]
    widths = [c[1] for c in PAYTRACK_COLS]
    tbl = doc.add_table(rows=1 + 12, cols=len(headers))
    tbl.style = "Table Grid"
    for i, txt in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(8)
    for r in range(1, 1 + 12):
        for i in range(len(headers)):
            for para in tbl.rows[r].cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    _set_widths(tbl, widths)
    doc.add_paragraph()
    _signature_block(doc, [
        "Period covered: ____________  to  ____________",
        "Total scripts approved this period: ______   Total gold seen: ______",
        "Total amount due this period: £______________",
        "Marker declaration: I confirm I genuinely reviewed every script "
        "recorded above in line with the Marker SOP and did not "
        "rubber-stamp.",
        "Marker signature: ____________________________   Date: ____________",
        "Company verification (reconciled vs system): ______________________",
    ])


def add_paid_marker_overview(doc):
    """Master-pack only: a short bridge explaining the paid-marker drive
    layered on top of the moderation drive, before the per-board sections."""
    doc.add_page_break()
    doc.add_heading("Paid-Marker Drive — engagement layer", level=1)
    _para(doc,
          "The moderation drive above is the mechanic. The paid-marker drive "
          "scales it: contracted, NDA'd English examiners are paid per "
          "approved script to genuinely review AI marking, with IP in their "
          "marking output assigned to the Company. Each board section that "
          "follows now also includes everything a marker needs:", italic=True)
    for s in [
        "Rubric & exemplar summary — calibrate to the mark-scheme intent.",
        "Marker SOP — exact console steps (log in → /marker → review AI "
        "draft → correct → fill adjustment reason → Approve & next).",
        "Source & consent matrix — commissioned/specimen need NO pupil "
        "consent; platform (b2c_self/b2b_class) needs student + parental "
        "consent.",
        "Calibration guide — gold scripts, target gold accuracy, drift "
        "consequences.",
        "Marker contract & NDA template — engagement, IP assignment, "
        "confidentiality, data handling, pay per approved script, quality.",
        "Payment & throughput tracker — per-session record reconciled "
        "against /admin/marker-pay.",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    _hr_box(doc, "Pay is for GENUINE review. Rubber-stamping is a contract "
                 "breach, corrupts the training signal, and is detected by "
                 "blind gold scripts and audit re-marks.")


def add_paid_marker_sections(doc, b):
    """All six paid-marker sections for one board, in reading order. Used by
    BOTH the per-board worksheet files and the master pack so each appears
    per board AND in the master."""
    doc.add_page_break()
    add_rubric_summary(doc, b)
    doc.add_page_break()
    add_marker_sop(doc, b)
    doc.add_page_break()
    add_source_consent_matrix(doc, b)
    doc.add_page_break()
    add_calibration_guide(doc, b)
    doc.add_page_break()
    add_contract_nda(doc, b)
    doc.add_page_break()
    add_payment_tracker(doc, b)


def main():
    # 1) Master pack: brief + every board section.
    master = Document()
    _landscape_narrow(master)
    build_brief(master)
    add_paid_marker_overview(master)
    for i, b in enumerate(BOARDS):
        master.add_page_break()
        add_board_section(master, b, first=True)
        # Per board, in the master: rubric & exemplars, SOP, source/consent
        # matrix, calibration guide, contract & NDA, payment tracker.
        add_paid_marker_sections(master, b)
    master_path = OUT / "00-Smart-IP-Moderation-Drive-Pack.docx"
    master.save(str(master_path))
    written = [master_path.name]

    # 2) One fillable worksheet per board (for hand-out) — now carries the
    #    paid-marker sections too so a marker has everything in one file.
    for b in BOARDS:
        d = Document()
        _landscape_narrow(d)
        add_board_section(d, b, first=True)
        add_paid_marker_sections(d, b)
        path = OUT / f"{b['file']}-Moderation-Worksheet.docx"
        d.save(str(path))
        written.append(path.name)

    print("WROTE:")
    for w in written:
        print("  business-docs/smart-ip/moderation-drive/" + w)


if __name__ == "__main__":
    main()
