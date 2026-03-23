"""
Convert all compliance markdown and HTML files to professional Word documents.
Output: D:/Coding/the-english-hub/Compliance folder/
"""

import os
import re
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


COMPANY_NAME = "Upskill Energy Limited"
BRAND_NAME = "The English Hub"
BRAND_COLOR = RGBColor(0x1A, 0x52, 0x76)  # Dark blue
ACCENT_COLOR = RGBColor(0x2E, 0x86, 0xC1)  # Medium blue
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)

SOURCE_DIR = "D:/Coding/the-english-hub/compliance"
OUTPUT_DIR = "D:/Coding/the-english-hub/Compliance folder"


def setup_styles(doc):
    """Configure professional document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = BRAND_COLOR
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = BRAND_COLOR
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ACCENT_COLOR
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = ACCENT_COLOR
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    return doc


def add_header_footer(doc, title):
    """Add professional header and footer."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"{BRAND_NAME}  |  {COMPANY_NAME}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("CONFIDENTIAL  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'
    run = fp.add_run(f"{COMPANY_NAME}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'


def add_cover_page(doc, title, category):
    """Add a professional cover page."""
    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(BRAND_NAME.upper())
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = 'Calibri'
    run.font.bold = True
    run.font.letter_spacing = Pt(3)

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = BRAND_COLOR
    run.font.size = Pt(12)

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_COLOR
    run.font.name = 'Calibri'
    run.font.bold = True

    # Category
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(category)
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = 'Calibri'

    # Spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Company details
    details = [
        f"Prepared for: {COMPANY_NAME}",
        "Platform: The English Hub",
        "Date: March 2026",
        "Classification: Confidential",
    ]
    for detail in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(detail)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = 'Calibri'

    # Page break after cover
    doc.add_page_break()


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_table(table):
    """Apply professional formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "1A5276")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

    # Style data rows with alternating colors
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, "F2F7FB")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def parse_markdown_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-\|:]+\|$', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows


def add_inline_formatting(paragraph, text):
    """Add text with bold/italic markdown formatting to a paragraph."""
    # Split by bold+italic, bold, italic patterns
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def process_markdown(doc, content):
    """Process markdown content and add to document."""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            # Check if next line is also a table line
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # End of table, process it
                in_table = False
                rows = parse_markdown_table(table_lines)
                if rows and len(rows) > 0:
                    max_cols = max(len(r) for r in rows)
                    if max_cols > 0:
                        try:
                            table = doc.add_table(rows=len(rows), cols=max_cols)
                            for r_idx, row_data in enumerate(rows):
                                for c_idx, cell_text in enumerate(row_data):
                                    if c_idx < max_cols:
                                        cell = table.cell(r_idx, c_idx)
                                        cell.text = ''
                                        p = cell.paragraphs[0]
                                        # Strip markdown bold from cell text
                                        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                                        p.add_run(clean)
                            format_table(table)
                            doc.add_paragraph()  # spacing after table
                        except Exception:
                            # Fallback: just add as text
                            for tl in table_lines:
                                doc.add_paragraph(tl)
                table_lines = []
                i += 1
                continue

        # Headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.*)', line)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                # Remove markdown formatting from heading text
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                text = re.sub(r'\*(.*?)\*', r'\1', text)
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                elif level == 3:
                    doc.add_heading(text, level=3)
                else:
                    doc.add_heading(text, level=4)
                i += 1
                continue

        # Horizontal rules
        if re.match(r'^[\-\*\_]{3,}\s*$', line.strip()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("━" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # Bullet points
        bullet_match = re.match(r'^(\s*)([\-\*\+])\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(3)
            level = min(indent // 2, 3)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
            add_inline_formatting(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\s*)(\d+[\.\)])\s+(.*)', line)
        if num_match:
            indent = len(num_match.group(1))
            text = num_match.group(3)
            level = min(indent // 3, 3)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
            add_inline_formatting(p, text)
            i += 1
            continue

        # Checkbox items
        checkbox_match = re.match(r'^(\s*)[\-\*]\s+\[([ xX~!])\]\s+(.*)', line)
        if checkbox_match:
            checked = checkbox_match.group(2).lower()
            text = checkbox_match.group(3)
            symbol = "☑" if checked in ('x',) else "☐" if checked == ' ' else "◻"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            run = p.add_run(f"{symbol}  ")
            run.font.size = Pt(12)
            add_inline_formatting(p, text)
            i += 1
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.right_indent = Cm(1.27)
            run = p.add_run("  " + text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraphs
        p = doc.add_paragraph()
        # Remove link markdown but keep text
        clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
        add_inline_formatting(p, clean_line)
        i += 1


def get_title_and_category(filepath, content):
    """Extract title from content and determine category."""
    # Get title from first heading
    match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if match:
        title = match.group(1).strip()
        title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
    else:
        title = Path(filepath).stem.replace('-', ' ').replace('_', ' ').title()

    # Determine category from path
    rel_path = filepath.replace(SOURCE_DIR + '/', '').replace(SOURCE_DIR + '\\', '')
    parts = rel_path.split('/')
    if len(parts) > 1:
        cat = parts[0].replace('-', ' ').replace('_', ' ').title()
    else:
        cat = "Compliance"

    category_map = {
        'Uk Legal': 'UK Legal Compliance',
        'Qatar Legal': 'Qatar Legal Compliance',
        'Cross Border': 'Cross-Border Compliance',
        'Policies': 'Company Policies',
        'Marketing': 'Marketing Compliance',
        'Technical': 'Technical Specifications',
        'Business': 'Business Operations',
        'Templates': 'Document Templates',
    }
    category = category_map.get(cat, cat)

    return title, category


def extract_text_from_html(filepath):
    """Extract readable text from HTML files, converting to markdown-like format."""
    with open(filepath, 'r', encoding='utf-8') as f:
        html_content = f.read()

    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Remove script and style elements
    for element in soup(['script', 'style']):
        element.decompose()

    # Get the title
    title_tag = soup.find('title')
    title = title_tag.get_text() if title_tag else Path(filepath).stem.replace('-', ' ').title()

    # Convert to text preserving some structure
    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append("*This document contains the technical specification and reference implementation for this component.*")
    lines.append("")

    # Process body content
    body = soup.find('body') or soup

    for element in body.children:
        if hasattr(element, 'name'):
            if element.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(element.name[1])
                lines.append(f"{'#' * level} {element.get_text().strip()}")
                lines.append("")
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    lines.append(text)
                    lines.append("")
            elif element.name in ('ul', 'ol'):
                for li in element.find_all('li', recursive=False):
                    lines.append(f"- {li.get_text().strip()}")
                lines.append("")
            elif element.name == 'div':
                text = element.get_text(separator='\n').strip()
                if text:
                    for subline in text.split('\n'):
                        subline = subline.strip()
                        if subline:
                            lines.append(subline)
                    lines.append("")
            elif element.name == 'table':
                rows = element.find_all('tr')
                for tr in rows:
                    cells = tr.find_all(['th', 'td'])
                    cell_texts = [c.get_text().strip() for c in cells]
                    lines.append("| " + " | ".join(cell_texts) + " |")
                lines.append("")

    # If we got very little content, fall back to full text extraction
    text = '\n'.join(lines)
    if len(text.strip()) < 200:
        full_text = soup.get_text(separator='\n')
        clean_lines = []
        for line in full_text.split('\n'):
            line = line.strip()
            if line:
                clean_lines.append(line)
        text = f"# {title}\n\n" + '\n\n'.join(clean_lines)

    return text


def convert_file(filepath, output_dir):
    """Convert a single file to a professional Word document."""
    filename = Path(filepath).stem
    ext = Path(filepath).suffix.lower()

    # Read content
    if ext == '.html':
        content = extract_text_from_html(filepath)
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

    # Get title and category
    title, category = get_title_and_category(filepath, content)

    # Create document
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc, title)
    add_cover_page(doc, title, category)

    # Process content (skip the first heading since it's on the cover page)
    # Remove the first # heading from content to avoid duplication
    content_lines = content.split('\n')
    skip_first_heading = False
    filtered_lines = []
    for line in content_lines:
        if not skip_first_heading and re.match(r'^#\s+', line):
            skip_first_heading = True
            continue
        filtered_lines.append(line)
    content = '\n'.join(filtered_lines)

    process_markdown(doc, content)

    # Generate output filename
    # Create a clean filename from the original
    clean_name = filename.replace('-', ' ').replace('_', ' ').title()
    output_path = os.path.join(output_dir, f"{clean_name}.docx")

    # Handle duplicates
    if os.path.exists(output_path):
        # Add category prefix
        cat_prefix = Path(filepath).parent.name.replace('-', ' ').title()
        output_path = os.path.join(output_dir, f"{cat_prefix} - {clean_name}.docx")

    doc.save(output_path)
    return output_path


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Collect all files
    files = []
    for ext in ('*.md', '*.html'):
        files.extend(glob.glob(os.path.join(SOURCE_DIR, '**', ext), recursive=True))

    files.sort()
    print(f"Found {len(files)} files to convert.\n")

    success = 0
    failed = []

    for filepath in files:
        filepath = filepath.replace('\\', '/')
        rel = filepath.replace(SOURCE_DIR + '/', '')
        try:
            output_path = convert_file(filepath, OUTPUT_DIR)
            out_name = Path(output_path).name
            print(f"  OK  {rel:60s} -> {out_name}")
            success += 1
        except Exception as e:
            print(f" FAIL {rel:60s} -> {e}")
            failed.append((rel, str(e)))

    print(f"\n{'='*70}")
    print(f"Converted: {success}/{len(files)} files")
    if failed:
        print(f"Failed: {len(failed)}")
        for f, err in failed:
            print(f"  - {f}: {err}")
    print(f"\nOutput directory: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
